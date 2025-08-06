import os
import json
import logging
import email
import re
import docx
import filetype
import pdfplumber
import PyPDF2
from pathlib import Path
from typing import List, Dict, Any, Optional, Tuple
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from openai import OpenAI, APIError
import io
import pandas as pd
import asyncio
from concurrent.futures import ThreadPoolExecutor, as_completed
import threading

# Try importing specialized libraries, gracefully handle if not installed
try:
    import camelot
    CAMELOT_INSTALLED = True
except ImportError:
    CAMELOT_INSTALLED = False
    logging.warning("Camelot not installed. Table extraction from PDFs may be less effective.")

try:
    from unstructured.partition.auto import partition
    UNSTRUCTURED_INSTALLED = True
except ImportError:
    UNSTRUCTURED_INSTALLED = False
    logging.warning("Unstructured library not installed. Advanced parsing for non-PDFs will be limited.")

try:
    from pdf2image import convert_from_path
    import pytesseract
    PYTESSERACT_INSTALLED = True
except ImportError:
    PYTESSERACT_INSTALLED = False
    logging.warning("pytesseract/pdf2image not installed. OCR for image-based PDFs will not work.")
    
# NOTE: Poppler must be installed separately for pdf2image/pytesseract to work.
# Instructions: https://github.com/oschwartz10612/poppler-windows/releases/

# Dummy config for demonstration, replace with your actual config
class Config:
    CHUNK_SIZE = 1200
    CHUNK_OVERLAP = 300
    OPENAI_API_KEY = "YOUR_API_KEY"
    DOCUMENTS_PATH = "./documents"

config = Config()

logger = logging.getLogger(__name__)

# --- Standalone helper functions for threading ---
# These functions now directly perform the I/O-bound task.

def _thread_process_pdf_page_with_pdfplumber(file_path: str, page_num: int) -> str:
    """Helper function to extract text from a single PDF page using pdfplumber."""
    try:
        with pdfplumber.open(file_path) as pdf:
            page = pdf.pages[page_num]
            page_text = page.extract_text()
            return page_text if page_text else ""
    except Exception as e:
        logger.warning(f"pdfplumber extraction failed for page {page_num} of {file_path}: {e}")
        return ""

def _thread_process_pdf_page_with_pypdf2(file_path: str, page_num: int) -> str:
    """Helper function to extract text from a single PDF page using PyPDF2."""
    try:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            page = reader.pages[page_num]
            page_text = page.extract_text()
            return page_text if page_text else ""
    except Exception as e:
        logger.warning(f"PyPDF2 extraction failed for page {page_num} of {file_path}: {e}")
        return ""

# --- DocumentLoader Class ---

class DocumentLoader:
    """
    A universal and highly robust DocumentLoader.
    This class is responsible for loading various document types, extracting content,
    and creating structured chunks for downstream processing.
    It leverages ThreadPoolExecutor for concurrent I/O operations to improve performance.
    """
    def __init__(self):
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=config.CHUNK_SIZE,
            chunk_overlap=config.CHUNK_OVERLAP,
            separators=["\n\n\n", "\n\n", "\n", ". ", "; ", ", ", " ", ""]
        )
        self.supported_formats = {'.pdf', '.docx', '.doc', '.eml', '.msg', '.txt', '.csv', '.html', '.json', '.xlsx'}
        self.llm_client = OpenAI(api_key=config.OPENAI_API_KEY)
        # Use a ThreadPoolExecutor for I/O-bound tasks
        self._executor = ThreadPoolExecutor(max_workers=os.cpu_count() or 6)

    def __del__(self):
        """Ensure the executor is shut down when the object is destroyed."""
        if hasattr(self, '_executor') and self._executor:
            self._executor.shutdown(wait=True) # wait=True to ensure all tasks are finished
    
    async def process_document_async(self, file_path: str) -> List[Document]:
        """
        Async wrapper for the main document processing method.
        This allows the FastAPI endpoint to not block while processing.
        """
        return await asyncio.to_thread(self.process_document, file_path)
    
    def process_document(self, file_path: str) -> List[Document]:
        """
        Main method to process a single document.
        It uses a fallback approach to extract text, then structures and chunks the content.
        """
        logger.info(f"Processing document: {file_path}")
        file_format = self._detect_format(file_path)
        
        text = ""
        tables_data = []

        try:
            if file_format == '.pdf':
                text, tables_data = self.load_pdf(file_path)
            elif file_format == '.docx':
                text = self.load_docx(file_path)
            elif file_format in ['.eml', '.msg']:
                text = self.load_email(file_path)
            else:
                text = self._unstructured_parse(file_path)
        except Exception as e:
            logger.error(f"Failed native parsing for {file_path}: {e}. Falling back to unstructured.")
            text = self._unstructured_parse(file_path)

        if not text.strip() and not tables_data:
            raise ValueError(f"No content extracted from {file_path}")

        document_name = Path(file_path).stem
        
        # Use LLM-based structuring only for smaller, prose-heavy documents
        prose_sections = self._llm_based_document_structuring(text) if text else []
        
        all_sections = prose_sections + tables_data
        
        return self.create_chunks(all_sections, document_name)

    def _detect_format(self, file_path: str) -> str:
        """Detects document format, falling back to extension if guessing fails."""
        try:
            kind = filetype.guess(file_path)
            return f".{kind.extension}" if kind else Path(file_path).suffix.lower()
        except Exception:
            return Path(file_path).suffix.lower()

    # --- Text Extraction Methods (Multiple Fallbacks) ---
    def load_pdf(self, file_path: str) -> Tuple[str, List[Dict[str, Any]]]:
        """
        Extracts text from PDF with threading for better performance.
        This method uses ThreadPoolExecutor to handle concurrent I/O.
        """
        text = ""
        tables_data = []
        
        try:
            with pdfplumber.open(file_path) as pdf:
                pages_count = len(pdf.pages)
            
            logger.info(f"Processing {pages_count} pages from {file_path} with threading.")
            
            # Prepare the arguments for the helper function
            pages_to_process = [(file_path, i) for i in range(pages_count)]
            
            # Use executor.map for a clean, parallel execution of the helper function
            texts = list(self._executor.map(_thread_process_pdf_page_with_pdfplumber, *zip(*pages_to_process)))
            
            text = "\n".join(texts)
            
            # Table extraction remains sequential
            tables_data = self._extract_tables_from_pdf(file_path)
            return self._clean_text(text), tables_data

        except Exception as e:
            logger.warning(f"Threaded pdfplumber extraction failed: {e}. Falling back to single-threaded PyPDF2.")
            try:
                with open(file_path, 'rb') as file:
                    reader = PyPDF2.PdfReader(file)
                    pages_count = len(reader.pages)
                
                # Use a sequential loop for the fallback
                texts = []
                for i in range(pages_count):
                    page_text = _thread_process_pdf_page_with_pypdf2(file_path, i)
                    texts.append(page_text)
                
                text = "\n".join(texts)
                return self._clean_text(text), tables_data
            except Exception as e2:
                logger.warning(f"PyPDF2 sequential extraction failed: {e2}. Falling back to OCR.")
                if PYTESSERACT_INSTALLED:
                    logger.info("Falling back to OCR.")
                    text = self._extract_with_ocr(file_path)
                return self._clean_text(text), tables_data
    
    def load_docx(self, file_path: str) -> str:
        doc = docx.Document(file_path)
        text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    text += "\n" + " | ".join(cells)
        return self._clean_text(text)

    def load_email(self, file_path: str) -> str:
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            msg = email.message_from_file(f)
        text = "\n".join(f"{h}: {msg.get(h)}" for h in ["Subject", "From", "To", "Date"] if msg.get(h)) + "\n\n"
        if msg.is_multipart():
            for part in msg.walk():
                if part.get_content_type() == "text/plain":
                    payload = part.get_payload(decode=True)
                    if payload:
                        text += payload.decode('utf-8', errors='ignore')
        else:
            payload = msg.get_payload(decode=True)
            if payload:
                text += payload.decode('utf-8', errors='ignore')
        return self._clean_text(text)

    def _unstructured_parse(self, file_path: str) -> str:
        if not UNSTRUCTURED_INSTALLED:
            return ""
        try:
            elements = partition(file_path=file_path)
            return self._clean_text("\n\n".join([str(e) for e in elements]))
        except Exception as e:
            logger.warning(f"Unstructured failed for {file_path}: {e}")
            return ""

    def _extract_with_ocr(self, file_path: str) -> str:
        if not PYTESSERACT_INSTALLED:
            return ""
        try:
            images = convert_from_path(file_path)
            return "\n".join([pytesseract.image_to_string(img) for img in images])
        except Exception as e:
            logger.warning(f"OCR failed for {file_path}: {e}")
            return ""

    def _clean_text(self, text: str) -> str:
        return ' '.join(text.replace('\x00', '').split())

    def _llm_based_document_structuring(self, text: str) -> List[Dict[str, Any]]:
        # For very large documents, use simpler chunking to avoid expensive LLM calls and context limits
        if len(text) > 40000:
            logger.info(f"Large document ({len(text)} chars), using enhanced chunking instead of LLM structuring")
            # Fallback to the reliable RecursiveCharacterTextSplitter for large documents
            return self._create_enhanced_chunks_from_text(text)
        
        # For smaller documents, use LLM structuring
        prompt = f"""
        You are a document parser. Break the following text into logical sections.
        For each section, provide a concise 'section_id', the full 'content', and a 'type' from [definition, benefit, exclusion, condition, general].
        Return only a JSON array of objects.
        Document Text:
        {text}
        JSON Response:
        """
        
        try:
            if not config.OPENAI_API_KEY or config.OPENAI_API_KEY == "YOUR_API_KEY":
                raise ValueError("API key not set or is a placeholder.")

            response = self.llm_client.chat.completions.create(
                model="gpt-4o-mini",
                messages=[{"role": "user", "content": prompt}],
                temperature=0.1,
                response_format={"type": "json_object"}
            )
            content = response.choices[0].message.content.strip()
            match = re.search(r'\[.*\]', content, re.DOTALL)
            if match:
                return json.loads(match.group(0))
            else:
                logger.warning("LLM response did not contain a valid JSON array. Trying to salvage.")
                try:
                    return json.loads(content)
                except json.JSONDecodeError:
                    raise ValueError("LLM response did not contain a valid JSON array.")
        except (APIError, json.JSONDecodeError, ValueError) as e:
            logger.warning(f"LLM-based structuring failed: {e}. Falling back to default chunking.")
            return [{"section_id": "FULL_DOCUMENT", "content": text, "type": "general"}]
    
    def _create_enhanced_chunks_from_text(self, text: str) -> List[Dict[str, Any]]:
        """
        Create enhanced chunks from large documents using a reliable text splitter.
        This function now correctly and reliably splits large text chunks.
        """
        chunks = self.text_splitter.split_text(text)
        sections = []
        for i, chunk in enumerate(chunks):
            sections.append({
                "section_id": f"CHUNK_{i+1}",
                "content": chunk,
                "type": "general"
            })
        return sections

    def _extract_tables_from_pdf(self, file_path: str) -> List[Dict[str, Any]]:
        tables = []
        try:
            with pdfplumber.open(file_path) as pdf:
                for page_num, page in enumerate(pdf.pages):
                    extracted_tables = page.extract_tables()
                    for i, table_data in enumerate(extracted_tables):
                        if not table_data or not table_data[0]:
                            continue
                            
                        headers_rows = [row for row in table_data if any(str(c).strip() for c in row)][:2]
                        data_start_row_index = len(headers_rows)
                        
                        if len(headers_rows) > 1:
                            first_row = [str(c).strip() if c else "" for c in headers_rows[0]]
                            second_row = [str(c).strip() if c else "" for c in headers_rows[1]]
                            
                            merged_headers = []
                            last_main_header = None
                            for j in range(len(first_row)):
                                current_main_header = first_row[j] if j < len(first_row) and first_row[j] else ''
                                current_sub_header = second_row[j] if j < len(second_row) and second_row[j] else ''
                                
                                if current_main_header:
                                    last_main_header = current_main_header
                                    if current_sub_header:
                                        merged_headers.append(f"{last_main_header} - {current_sub_header}")
                                    else:
                                        merged_headers.append(last_main_header)
                                elif current_sub_header:
                                    merged_headers.append(f"{last_main_header} - {current_sub_header}")
                                else:
                                    merged_headers.append(f"Col_{len(merged_headers)}")
                            
                            headers = merged_headers
                            data_start_row_index = len(headers_rows)
                        else:
                            headers = [str(cell).strip() if cell else f"Col_{j}" for j, cell in enumerate(headers_rows[0])]
                            data_start_row_index = 1
                            
                        data_rows = table_data[data_start_row_index:]
                        
                        if not headers or not data_rows:
                            continue
                        
                        rows_list = []
                        for row in data_rows:
                            row_dict = {}
                            for j, header in enumerate(headers):
                                cell_value = str(row[j]).strip() if j < len(row) and row[j] else ""
                                row_dict[header] = cell_value
                            
                            if any(v for v in row_dict.values()):
                                rows_list.append(row_dict)

                        if rows_list:
                            tables.append({
                                'section_id': f'table_{page_num+1}_{i+1}',
                                'content': rows_list,
                                'type': 'table'
                            })
        except Exception as e:
            logger.warning(f"pdfplumber table extraction failed for {file_path}: {e}")
        return tables

    def create_chunks(self, sections: List[Dict[str, Any]], document_name: str) -> List[Document]:
        documents = []
        for section in sections:
            section_id = section.get('section_id', 'unknown')
            section_type = section.get('type', 'general')

            if section_type == 'table' and 'content' in section and isinstance(section['content'], list):
                table_rows = section['content']
                for i, row_data in enumerate(table_rows):
                    if not row_data or not isinstance(row_data, dict):
                        continue
                    
                    primary_key = None
                    primary_value = None
                    
                    for key, value in row_data.items():
                        if value and str(value).strip():
                            primary_key = key
                            primary_value = str(value).strip()
                            break
                    
                    if not primary_key or not primary_value:
                        continue
                    
                    data_pairs = []
                    for key, value in row_data.items():
                        if key != primary_key and value and str(value).strip():
                            clean_key = self._clean_column_name(key)
                            clean_value = str(value).strip()
                            if clean_value.lower() not in ['', 'na', 'n/a', 'null', 'none']:
                                data_pairs.append((clean_key, clean_value))
                    
                    if not data_pairs:
                        continue
                    
                    clean_primary = self._clean_text_content(primary_value)
                    if not clean_primary:
                        continue
                    
                    # Corrected logic: Generate prose from table data for hashing
                    row_content = self._generate_table_prose(clean_primary, data_pairs)
                    
                    if not row_content:
                        continue

                    documents.append(Document(page_content=row_content, metadata={
                        'document_name': document_name,
                        'section_id': section_id,
                        'section_type': 'table_row',
                        'chunk_index': i,
                        'source': f"{document_name}_{section_id}_row_{i}",
                        'chunk_strategy': 'table_prose',
                        'row_data': json.dumps(row_data),
                        'primary_key': primary_key,
                        'data_columns': len(data_pairs)
                    }))
            else:
                content = section.get('content', '').strip()
                if not content:
                    continue
                
                chunks = self._smart_chunk_with_definitions(content)

                for i, chunk in enumerate(chunks):
                    documents.append(Document(page_content=chunk, metadata={
                        'document_name': document_name,
                        'section_id': section_id,
                        'section_type': section_type,
                        'chunk_index': i,
                        'source': f"{document_name}_{section_id}_{i}",
                        'chunk_strategy': 'definition_aware'
                    }))
        return documents

    def _smart_chunk_with_definitions(self, content: str) -> List[str]:
        """Smart chunking that preserves complete policy definitions and key facts."""
        if not content.strip():
            return []
        
        definition_patterns = [
            r'(\d+\.\d+\s+[A-Z][^\n]*?\s+means\s+[^\n]*?(?:\n[^\n]*?)*?)(?=\n\n|\n\d+\.\d+|$)',
            r'([A-Z][^\n]*?\s+means\s+[^\n]*?(?:\n[^\n]*?)*?)(?=\n\n|\n[A-Z][^\n]*?\s+means|$)',
            r'(\"[^\"]+\"\s+means\s+[^\n]*?(?:\n[^\n]*?)*?)(?=\n\n|\n\"[^\"]+\"|$)',
            r'([A-Z][A-Z\s]+:\s*[^\n]*?(?:\n[^\n]*?)*?)(?=\n\n|\n[A-Z][A-Z\s]+:|$)'
        ]
        
        preserved_definitions = []
        remaining_content = content
        
        for pattern in definition_patterns:
            matches = re.finditer(pattern, remaining_content, re.MULTILINE | re.DOTALL)
            for match in matches:
                definition_text = match.group(1).strip()
                if len(definition_text) > 50:
                    preserved_definitions.append(definition_text)
                    remaining_content = remaining_content.replace(match.group(0), '', 1)
        
        remaining_content = re.sub(r'\n\s*\n\s*\n', '\n\n', remaining_content).strip()
        
        regular_chunks = self.text_splitter.split_text(remaining_content) if remaining_content else []
        
        all_chunks = preserved_definitions + regular_chunks
        
        final_chunks = [chunk.strip() for chunk in all_chunks if chunk.strip() and len(chunk.strip()) > 20]
        
        if not final_chunks:
            final_chunks = self.text_splitter.split_text(content)
        
        return final_chunks

    def _clean_column_name(self, column_name: str) -> str:
        """Clean and normalize column names for better readability."""
        if not column_name:
            return ''
        clean_name = str(column_name).strip()
        clean_name = re.sub(r'[*#@$%^&()\[\]{}|\\:;"<>?/~`]', '', clean_name)
        clean_name = re.sub(r'\s+', ' ', clean_name).strip()
        clean_name = re.sub(r'^(Plans?\s*-\s*|Plan\s*-\s*)', '', clean_name, flags=re.IGNORECASE)
        clean_name = re.sub(r'\s*(Column|Col)\s*\d+', '', clean_name, flags=re.IGNORECASE)
        return clean_name.strip()
    
    def _clean_text_content(self, text: str) -> str:
        """Clean and normalize text content for better readability."""
        if not text:
            return ''
        clean_text = str(text).strip()
        clean_text = re.sub(r'[*#@$%^&()\[\]{}|\\:;"<>?/~`]+', '', clean_text)
        clean_text = re.sub(r'\s+', ' ', clean_text).strip()
        clean_text = re.sub(r'\s*\([^)]*\)\s*', ' ', clean_text).strip()
        return clean_text
    
    def _generate_table_prose(self, primary_item: str, data_pairs: List[tuple]) -> str:
        """Generate natural language prose from table row data."""
        if not primary_item or not data_pairs:
            return ''
        
        values = [pair[1] for pair in data_pairs]
        unique_values = list(set(values))
        
        if len(unique_values) == 1:
            common_value = unique_values[0]
            if len(data_pairs) == 1:
                return f"{primary_item}: {common_value}."
            else:
                column_names = [pair[0] for pair in data_pairs]
                if len(column_names) <= 3:
                    columns_text = ', '.join(column_names)
                    return f"{primary_item} ({columns_text}): {common_value}."
                else:
                    return f"{primary_item}: {common_value} across all categories."
        else:
            details = []
            for column_name, value in data_pairs:
                if column_name and value:
                    details.append(f"{column_name}: {value}")
            
            if len(details) <= 3:
                return f"{primary_item} - " + "; ".join(details) + "."
            else:
                return f"{primary_item} has varying specifications: " + "; ".join(details[:3]) + f" and {len(details)-3} more."

    def load_all_documents(self, directory: str = None) -> List[Document]:
        """
        Loads all supported documents from a directory and processes them in a single batch.
        This method will now leverage parallel processing for PDFs.
        """
        doc_dir = Path(directory or config.DOCUMENTS_PATH)
        if not doc_dir.is_dir():
            logger.error(f"Directory not found: {doc_dir}")
            return []

        all_files = [str(file) for file in doc_dir.iterdir() if file.suffix.lower() in self.supported_formats]
        if not all_files:
            logger.warning(f"No supported documents found in {doc_dir}")
            return []

        all_documents = []
        for file_path in all_files:
            try:
                # The process_document call will now handle parallelization internally for PDFs
                processed_docs = self.process_document(file_path)
                all_documents.extend(processed_docs)
            except Exception as e:
                logger.error(f"Skipping document due to error: {file_path} - {e}")
                
        logger.info(f"Finished processing all documents. Total chunks created: {len(all_documents)}")
        return all_documents